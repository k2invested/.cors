#!/usr/bin/env python3
"""doc_edit_batch — batch edit a .docx file by paragraph line numbers.

Line numbers correspond to the output of doc_read.py (1-indexed, one per paragraph).
The LLM reads the numbered paragraph view, then specifies edits against those numbers.

Input JSON:
{
  "path": "<relative path to .docx file>",
  "edits": [
    {"lines": [15, 18], "action": "remove"},
    {"lines": [5, 5], "action": "replace", "content": "New paragraph text..."},
    {"lines": [22, 22], "action": "replace", "content": "Rewritten content..."},
    {"action": "replace_text", "find": "old phrase", "replace": "new phrase"}
  ]
}

Actions:
  remove — delete paragraphs in the line range (inclusive)
  replace — replace paragraphs in the line range with new content (one paragraph)
  replace_text — find/replace text across all paragraphs (no line range needed)

Env: WORKSPACE — sandbox root.
"""
import json
import os
import sys
from copy import deepcopy

sys.path.insert(0, os.path.dirname(__file__))
from scan_tree import sandbox_path

MAX_OUTPUT = 32_000


def load_docx(filepath):
    """Load a .docx file and return the Document object."""
    from docx import Document
    return Document(filepath)


def apply_edits(doc, edits):
    """Apply edits to the document. Returns list of change descriptions."""
    paragraphs = list(doc.paragraphs)
    total = len(paragraphs)
    changes = []

    # Collect line ranges to remove (process in reverse to preserve indices)
    removals = []
    replacements = []
    text_replacements = []

    for edit in edits:
        action = edit.get("action", "")

        if action == "replace_text":
            find = edit.get("find", "")
            replace = edit.get("replace", "")
            if find:
                text_replacements.append((find, replace))
            continue

        lines = edit.get("lines", [])
        if not lines or len(lines) != 2:
            changes.append(f"SKIPPED: invalid lines {lines}")
            continue

        start, end = lines[0], lines[1]
        # Convert 1-indexed to 0-indexed
        start_idx = start - 1
        end_idx = end - 1

        if start_idx < 0 or end_idx >= total or start_idx > end_idx:
            changes.append(f"SKIPPED: lines [{start}, {end}] out of range (1-{total})")
            continue

        if action == "remove":
            removals.append((start_idx, end_idx, start, end))
        elif action == "replace":
            content = edit.get("content", "")
            replacements.append((start_idx, end_idx, content, start, end))
        else:
            changes.append(f"SKIPPED: unknown action '{action}'")

    # Apply text replacements first (non-destructive)
    for find, replace in text_replacements:
        count = 0
        for para in paragraphs:
            for run in para.runs:
                if find in run.text:
                    run.text = run.text.replace(find, replace)
                    count += 1
        changes.append(f"replace_text: '{find}' → '{replace}' ({count} occurrences)")

    # Apply replacements (set text of first paragraph, mark rest for removal)
    for start_idx, end_idx, content, start, end in replacements:
        # Clear all runs in the first paragraph, add new text
        para = paragraphs[start_idx]
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = content
        else:
            para.add_run(content)

        # Mark remaining paragraphs in range for removal
        for idx in range(start_idx + 1, end_idx + 1):
            removals.append((idx, idx, idx + 1, idx + 1))

        preview = content[:60] + "..." if len(content) > 60 else content
        changes.append(f"replace lines [{start}-{end}]: '{preview}'")

    # Apply removals in reverse order to preserve indices
    # Deduplicate and sort
    remove_indices = set()
    for start_idx, end_idx, _, _ in removals:
        for idx in range(start_idx, end_idx + 1):
            remove_indices.add(idx)

    for idx in sorted(remove_indices, reverse=True):
        para = paragraphs[idx]
        # Remove paragraph from document by clearing its XML element
        p_element = para._element
        p_element.getparent().remove(p_element)

    if remove_indices:
        # Group consecutive indices for readable output
        sorted_indices = sorted(remove_indices)
        ranges = []
        range_start = sorted_indices[0] + 1  # back to 1-indexed
        prev = sorted_indices[0]
        for idx in sorted_indices[1:]:
            if idx == prev + 1:
                prev = idx
            else:
                range_end = prev + 1
                ranges.append(f"{range_start}-{range_end}" if range_start != range_end else str(range_start))
                range_start = idx + 1
                prev = idx
        range_end = prev + 1
        ranges.append(f"{range_start}-{range_end}" if range_start != range_end else str(range_start))
        changes.append(f"removed lines: {', '.join(ranges)} ({len(remove_indices)} paragraphs)")

    return changes


def main():
    params = json.load(sys.stdin)
    workspace = os.environ.get("WORKSPACE", ".")

    file_path = params.get("path", "")
    edits = params.get("edits", [])

    if not file_path:
        print("Error: missing 'path' parameter", file=sys.stderr)
        sys.exit(1)

    if not edits:
        print("Error: missing 'edits' parameter", file=sys.stderr)
        sys.exit(1)

    try:
        resolved = sandbox_path(file_path, workspace)
    except (ValueError, FileNotFoundError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(resolved):
        print(f"Error: '{file_path}' is not a file", file=sys.stderr)
        sys.exit(1)

    # Load document
    try:
        doc = load_docx(resolved)
    except Exception as e:
        print(f"Error loading docx: {e}", file=sys.stderr)
        sys.exit(1)

    # Count paragraphs before
    before_count = len(doc.paragraphs)

    # Apply edits
    changes = apply_edits(doc, edits)

    # Save
    try:
        doc.save(resolved)
    except Exception as e:
        print(f"Error saving docx: {e}", file=sys.stderr)
        sys.exit(1)

    # Report
    after_doc = load_docx(resolved)
    after_count = len(after_doc.paragraphs)

    output_parts = [f"Updated {file_path} ({before_count} → {after_count} paragraphs)"]
    for change in changes:
        output_parts.append(f"  {change}")

    output = "\n".join(output_parts)
    if len(output) > MAX_OUTPUT:
        output = output[:MAX_OUTPUT] + "\n... [truncated]"
    print(output)


if __name__ == "__main__":
    main()
