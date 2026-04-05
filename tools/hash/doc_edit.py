#!/usr/bin/env python3
"""doc_edit — modify document files (.docx) in place.

Input JSON: {
  "path": "<relative file path>",
  "op": "rewrite | replace_text | add_paragraph | replace_paragraph | set_property",
  ...
}
Env: WORKSPACE — sandbox root.

Operations:
  rewrite            — rebuild the entire document from a list of paragraphs
  replace_text       — find/replace text across all paragraphs (exact match)
  add_paragraph      — append a new paragraph at the end
  replace_paragraph  — replace paragraph at index with new text
  set_property       — set document property (title, author, subject, keywords)
"""
import json, os, sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
from tools.scan_tree import sandbox_path


def main():
    params = json.load(sys.stdin)
    workspace = os.environ.get("WORKSPACE", ".")

    file_path = params.get("path", "")
    op = params.get("op", "")

    if not file_path:
        print("Error: missing 'path' parameter", file=sys.stderr)
        sys.exit(1)
    if not op:
        print("Error: missing 'op' parameter", file=sys.stderr)
        sys.exit(1)

    try:
        resolved = sandbox_path(file_path, workspace)
    except (ValueError, FileNotFoundError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(resolved):
        print(f"Error: '{file_path}' is not a file", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(resolved)[1].lower()
    if ext != ".docx":
        print(f"Error: unsupported format '{ext}'. Supported: .docx", file=sys.stderr)
        sys.exit(1)

    from docx import Document

    if op == "rewrite":
        # Rebuild entire document from a list of paragraphs
        paragraphs = params.get("paragraphs", [])
        if not paragraphs:
            print("Error: 'paragraphs' list is required for rewrite", file=sys.stderr)
            sys.exit(1)

        doc = Document()
        for p in paragraphs:
            if isinstance(p, str):
                doc.add_paragraph(p)
            elif isinstance(p, dict):
                text = p.get("text", "")
                style = p.get("style", None)
                if style:
                    try:
                        doc.add_paragraph(text, style=style)
                    except KeyError:
                        doc.add_paragraph(text)
                else:
                    doc.add_paragraph(text)
        doc.save(resolved)
        print(f"ok: {file_path} — rewritten with {len(paragraphs)} paragraphs")
        return

    doc = Document(resolved)

    if op == "replace_text":
        find = params.get("find", "")
        replace = params.get("replace", "")
        if not find:
            print("Error: 'find' is required for replace_text", file=sys.stderr)
            sys.exit(1)

        count = 0
        for para in doc.paragraphs:
            if find in para.text:
                # Preserve runs structure where possible
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1
                # Fallback: if find spans multiple runs, do full paragraph replace
                if count == 0 and find in para.text:
                    full = para.text.replace(find, replace)
                    for i, run in enumerate(para.runs):
                        if i == 0:
                            run.text = full
                        else:
                            run.text = ""
                    count += 1

        if count == 0:
            print(f"Warning: text '{find[:80]}' not found in any paragraph")
            sys.exit(0)
        doc.save(resolved)
        print(f"ok: {file_path} — replaced '{find[:50]}' → '{replace[:50]}' ({count} occurrence(s))")

    elif op == "add_paragraph":
        text = params.get("text", "")
        style = params.get("style", None)
        if not text:
            print("Error: 'text' is required for add_paragraph", file=sys.stderr)
            sys.exit(1)
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
        doc.save(resolved)
        print(f"ok: {file_path} — added paragraph ({len(text)} chars)")

    elif op == "replace_paragraph":
        index = params.get("index")
        text = params.get("text", "")
        style = params.get("style", None)
        if index is None:
            print("Error: 'index' is required for replace_paragraph", file=sys.stderr)
            sys.exit(1)
        index = int(index)
        if index < 0 or index >= len(doc.paragraphs):
            print(f"Error: index {index} out of range (0-{len(doc.paragraphs)-1})", file=sys.stderr)
            sys.exit(1)
        para = doc.paragraphs[index]
        # Clear existing runs
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)
        if style:
            para.style = doc.styles[style]
        doc.save(resolved)
        print(f"ok: {file_path} — replaced paragraph {index}")

    elif op == "set_property":
        prop = params.get("property", "")
        value = params.get("value", "")
        if not prop:
            print("Error: 'property' is required for set_property", file=sys.stderr)
            sys.exit(1)
        core = doc.core_properties
        if prop == "title":
            core.title = value
        elif prop == "author":
            core.author = value
        elif prop == "subject":
            core.subject = value
        elif prop == "keywords":
            core.keywords = value
        else:
            print(f"Error: unknown property '{prop}'. Supported: title, author, subject, keywords", file=sys.stderr)
            sys.exit(1)
        doc.save(resolved)
        print(f"ok: {file_path} — set {prop} = '{value[:50]}'")

    else:
        print(f"Error: unknown op '{op}'. Must be: replace_text, add_paragraph, replace_paragraph, set_property", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
