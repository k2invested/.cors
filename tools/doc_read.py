#!/usr/bin/env python3
"""doc_read — extract text content from document files (.docx).

Input JSON: {"path": "<relative path to document file>"}
Env: WORKSPACE — sandbox root.

Reads .docx files using python-docx, returns paragraph text with structure.
"""
import json, os, sys

sys.path.insert(0, os.path.dirname(__file__))
from scan_tree import sandbox_path

MAX_OUTPUT = 32_000


def read_docx(filepath):
    """Extract text from a .docx file with stable paragraph numbering.

    Every paragraph gets a line number (1-indexed), including empty ones.
    These numbers are stable references for doc_edit_batch.py edits.
    """
    from docx import Document
    doc = Document(filepath)

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if "Heading" in style and text:
            level = style.replace("Heading ", "").strip()
            lines.append(f"[h{level}] {text}")
        else:
            lines.append(text)

    # Append tables at the end
    for t_idx, table in enumerate(doc.tables):
        lines.append(f"[table {t_idx + 1}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def main():
    params = json.load(sys.stdin)
    workspace = os.environ.get("WORKSPACE", ".")

    file_path = params.get("path", "")
    if not file_path:
        print("Error: missing 'path' parameter", file=sys.stderr)
        sys.exit(1)

    try:
        resolved = sandbox_path(file_path, workspace)
    except (ValueError, FileNotFoundError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(resolved):
        print(f"Error: '{file_path}' is not a file", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(resolved)[1].lower()

    if ext == ".docx":
        try:
            content = read_docx(resolved)
        except Exception as e:
            print(f"Error reading docx: {e}", file=sys.stderr)
            sys.exit(1)
    else:
        print(f"Error: unsupported document format '{ext}'. Supported: .docx", file=sys.stderr)
        sys.exit(1)

    if not content.strip():
        print(f"(document is empty: {file_path})")
    else:
        # Output with line numbers (matches scan_tree format for batch compatibility)
        lines = content.split("\n")
        numbered = [f"{i+1}: {line}" for i, line in enumerate(lines)]
        output = "\n".join(numbered)
        if len(output) > MAX_OUTPUT:
            output = output[:MAX_OUTPUT] + f"\n... (truncated, {len(lines)} lines total)"
        print(output)


if __name__ == "__main__":
    main()
